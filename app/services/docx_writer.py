"""Build an annotated .docx from the extracted script + findings.

Approach:
1. Use python-docx to lay out the body with commentRangeStart/End + commentReference markers.
2. After saving, re-open the .docx as a ZIP to add `word/comments.xml`, register
   its content type, and add the document-level relationship. This is the
   supported pattern for Word comments since python-docx has no high-level API.

Result opens with native comments in Word, Google Docs, and LibreOffice.
"""
from __future__ import annotations

import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)

SEV_LABEL = {"high": "HIGH", "medium": "MED", "low": "LOW"}
HIGHLIGHT = {"high": "FFC7CE", "medium": "FFEB9C", "low": "D9D9D9"}


def build_annotated_docx(
    script_text: str,
    findings: list[dict],
    out_path: str,
    source_filename: str = "script",
) -> str:
    doc = Document()
    doc.add_heading(f"S&P Review — {source_filename}", level=1)
    intro = doc.add_paragraph(
        f"{len(findings)} finding(s). Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}."
    )
    for r in intro.runs:
        r.italic = True

    lines = script_text.splitlines()
    flags_by_line: dict[int, list[dict]] = {}
    for idx, f in enumerate(findings):
        for ln in range(f["line_start"], f["line_end"] + 1):
            flags_by_line.setdefault(ln, []).append({**f, "_id": idx})

    for i, line in enumerate(lines, start=1):
        para = doc.add_paragraph()
        flags_here = flags_by_line.get(i, [])
        starts = [f for f in flags_here if f["line_start"] == i]
        ends = [f for f in flags_here if f["line_end"] == i]

        for f in starts:
            _append_range_start(para, f["_id"])
        run = para.add_run(line if line else " ")
        if flags_here:
            _highlight_run(run, flags_here[0]["severity"])
        for f in ends:
            _append_range_end(para, f["_id"])
            _append_comment_reference(para, f["_id"])

    doc.save(out_path)

    if findings:
        _inject_comments_zip(out_path, findings)
    return out_path


def _highlight_run(run, sev: str) -> None:
    color = HIGHLIGHT.get(sev, "FFEB9C")
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    rPr.append(shd)


def _append_range_start(para, cid: int) -> None:
    el = OxmlElement("w:commentRangeStart")
    el.set(qn("w:id"), str(cid))
    para._p.append(el)


def _append_range_end(para, cid: int) -> None:
    el = OxmlElement("w:commentRangeEnd")
    el.set(qn("w:id"), str(cid))
    para._p.append(el)


def _append_comment_reference(para, cid: int) -> None:
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "CommentReference")
    rPr.append(rStyle)
    run.append(rPr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(cid))
    run.append(ref)
    para._p.append(run)


def _build_comments_xml(findings: list[dict]) -> bytes:
    root = etree.Element(f"{{{W_NS}}}comments", nsmap={"w": W_NS})
    for cid, f in enumerate(findings):
        c = etree.SubElement(root, f"{{{W_NS}}}comment")
        c.set(f"{{{W_NS}}}id", str(cid))
        c.set(f"{{{W_NS}}}author", "SnP Checker")
        c.set(f"{{{W_NS}}}initials", "SNP")
        c.set(f"{{{W_NS}}}date", datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"))
        body_lines = _comment_body(f).splitlines()
        p = etree.SubElement(c, f"{{{W_NS}}}p")
        for idx, line in enumerate(body_lines):
            r = etree.SubElement(p, f"{{{W_NS}}}r")
            t = etree.SubElement(r, f"{{{W_NS}}}t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = line
            if idx < len(body_lines) - 1:
                etree.SubElement(r, f"{{{W_NS}}}br")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _comment_body(f: dict) -> str:
    sev = SEV_LABEL.get(f.get("severity", ""), f.get("severity", ""))
    parts = [f"[{sev}] {f.get('reason', '')}"]
    if f.get("guideline_ref"):
        parts.append(f"Guideline: {f['guideline_ref']}")
    if f.get("suggestion"):
        parts.append(f"Fix: {f['suggestion']}")
    return "\n".join(parts)


def _inject_comments_zip(docx_path: str, findings: list[dict]) -> None:
    """Add comments.xml + content type + relationship into the .docx zip."""
    src = Path(docx_path)
    tmp = src.with_suffix(".tmp.docx")

    comments_xml = _build_comments_xml(findings)

    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        names = set(zin.namelist())
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = _patch_content_types(data)
            elif item.filename == "word/_rels/document.xml.rels":
                data = _patch_document_rels(data)
            zout.writestr(item, data)
        zout.writestr("word/comments.xml", comments_xml)
        if "word/_rels/document.xml.rels" not in names:
            zout.writestr("word/_rels/document.xml.rels", _new_document_rels())

    shutil.move(tmp, src)


def _patch_content_types(xml: bytes) -> bytes:
    ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    root = etree.fromstring(xml)
    for child in root.findall(f"{{{ns}}}Override"):
        if child.get("PartName") == "/word/comments.xml":
            return xml  # already present
    override = etree.SubElement(root, f"{{{ns}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set("ContentType", COMMENTS_CT)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _patch_document_rels(xml: bytes) -> bytes:
    ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    root = etree.fromstring(xml)
    for rel in root.findall(f"{{{ns}}}Relationship"):
        if rel.get("Target") == "comments.xml":
            return xml
    existing_ids = {rel.get("Id") for rel in root.findall(f"{{{ns}}}Relationship")}
    new_id = _next_rid(existing_ids)
    rel = etree.SubElement(root, f"{{{ns}}}Relationship")
    rel.set("Id", new_id)
    rel.set("Type", COMMENTS_REL_TYPE)
    rel.set("Target", "comments.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _new_document_rels() -> bytes:
    ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    root = etree.Element(f"{{{ns}}}Relationships")
    rel = etree.SubElement(root, f"{{{ns}}}Relationship")
    rel.set("Id", "rId1")
    rel.set("Type", COMMENTS_REL_TYPE)
    rel.set("Target", "comments.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _next_rid(existing: set[str]) -> str:
    n = 1
    while f"rId{n}" in existing:
        n += 1
    return f"rId{n}"
